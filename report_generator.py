from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def clean(text):
    if not text:
        return ""
    return text.encode('ascii', 'ignore').decode('ascii')


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_cell_text(cell, text, bold=False, font_size=10, color=None):
    para = cell.paragraphs[0]
    run = para.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, size=14, color="1F3864", bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(clean(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_report(results: list[dict]) -> bytes:
    today = datetime.now().strftime("%B %d, %Y")
    top_picks = [r for r in results if r["relevance_score"] >= 7][:5]

    sector_counts = {}
    for r in results:
        s = r.get("gfam_sector", "Other")
        sector_counts[s] = sector_counts.get(s, 0) + 1

    sector_summaries = {}
    for sector in ["Infrastructure", "Healthcare Services", "Financial Services", "Special Situations"]:
        sector_articles = [r for r in results if r.get("gfam_sector") == sector]
        if sector_articles:
            sector_summaries[sector] = max(sector_articles, key=lambda x: x["relevance_score"])

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Header ───────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)

    r1 = p.add_run("GFAM DEAL FLOW  |  DAILY BRIEFING")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor.from_string("1F3864")

    r2 = p.add_run(f"    {clean(today)}")
    r2.font.size = Pt(10)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor.from_string("888888")

    doc.add_paragraph()

    # ── Market Overview ───────────────────────────────────────
    add_heading(doc, "MARKET OVERVIEW", size=11)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3.5)

    row = table.add_row()
    set_cell_bg(row.cells[0], "DCE6F1")
    set_cell_bg(row.cells[1], "F5F5F5")
    add_cell_text(row.cells[0], "Articles Analyzed", bold=True)
    add_cell_text(row.cells[1], str(len(results)))

    for sector in ["Infrastructure", "Healthcare Services", "Financial Services", "Special Situations"]:
        count = sector_counts.get(sector, 0)
        if count > 0:
            row = table.add_row()
            set_cell_bg(row.cells[0], "EEF3FB")
            add_cell_text(row.cells[0], sector, bold=True)
            add_cell_text(row.cells[1], f"{count} article{'s' if count > 1 else ''}")

    doc.add_paragraph()
    add_divider(doc)

    # ── Top Opportunities ─────────────────────────────────────
    add_heading(doc, "TOP OPPORTUNITIES", size=12)

    if not top_picks:
        p = doc.add_paragraph("No articles scored 7 or above.")
        p.runs[0].font.name = "Arial"
        p.runs[0].font.size = Pt(10)
    else:
        for idx, item in enumerate(top_picks):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(f"{idx+1}. {clean(item['title'])}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Arial"

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            meta = f"{clean(item.get('gfam_sector',''))}  |  {clean(item.get('deal_type',''))}  |  Score: {item['relevance_score']}/10  |  {clean(item.get('source',''))}"
            r2 = p2.add_run(meta)
            r2.font.size = Pt(9)
            r2.font.name = "Arial"
            r2.font.color.rgb = RGBColor.from_string("888888")

            cap_fit = clean(item.get("gfam_capital_fit", ""))
            if cap_fit and cap_fit != "None":
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_after = Pt(2)
                r3a = p3.add_run("Capital fit: ")
                r3a.bold = True
                r3a.font.size = Pt(10)
                r3a.font.name = "Arial"
                r3b = p3.add_run(cap_fit)
                r3b.font.size = Pt(10)
                r3b.font.name = "Arial"
                r3b.font.color.rgb = RGBColor.from_string("1F5C2E")

            rationale = clean(item.get("investment_rationale", ""))
            if rationale:
                p4 = doc.add_paragraph()
                p4.paragraph_format.space_after = Pt(8)
                r4 = p4.add_run(rationale)
                r4.italic = True
                r4.font.size = Pt(10)
                r4.font.name = "Arial"
                r4.font.color.rgb = RGBColor.from_string("444444")

            if idx < len(top_picks) - 1:
                add_divider(doc)

    doc.add_paragraph()
    add_divider(doc)

    # ── Sector Highlights ─────────────────────────────────────
    add_heading(doc, "SECTOR HIGHLIGHTS", size=12)
    doc.add_paragraph()

    stable = doc.add_table(rows=0, cols=3)
    stable.style = 'Table Grid'
    stable.autofit = False
    stable.columns[0].width = Inches(1.5)
    stable.columns[1].width = Inches(2.9)
    stable.columns[2].width = Inches(2.1)

    hrow = stable.add_row()
    for cell, label in zip(hrow.cells, ["Sector", "Best Article", "Investment Angle"]):
        set_cell_bg(cell, "DCE6F1")
        add_cell_text(cell, label, bold=True, font_size=9)

    for sector in ["Infrastructure", "Healthcare Services", "Financial Services", "Special Situations"]:
        if sector_counts.get(sector, 0) == 0:
            continue
        best = sector_summaries.get(sector, {})
        row = stable.add_row()
        add_cell_text(row.cells[0], sector, bold=True, font_size=9)
        add_cell_text(row.cells[1], clean(best.get("title", "-")), font_size=9)
        rationale = clean(best.get("investment_rationale", "-"))[:130]
        add_cell_text(row.cells[2], rationale, font_size=9)

    # ── Footer ────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fr = fp.add_run(f"GFAM Deal Flow Monitor  |  Confidential  |  {clean(today)}")
    fr.font.size = Pt(8)
    fr.font.name = "Arial"
    fr.font.color.rgb = RGBColor.from_string("999999")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
